"""
Format the final transcript into a clean, readable docx.
- Bold speaker names
- Timestamps in gray
- Long segments split into natural paragraphs at sentence boundaries
- Visual separation between different speakers
"""

import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_ts(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def split_into_paragraphs(text, max_chars=400):
    """Split long text into paragraphs at sentence boundaries."""
    if len(text) <= max_chars:
        return [text]

    sentences = re.split(r'(?<=[.!?])\s+', text)
    paragraphs = []
    current = []
    current_len = 0

    for sentence in sentences:
        if current_len + len(sentence) > max_chars and current:
            paragraphs.append(' '.join(current))
            current = [sentence]
            current_len = len(sentence)
        else:
            current.append(sentence)
            current_len += len(sentence) + 1

    if current:
        paragraphs.append(' '.join(current))

    return paragraphs


def clean_text(text):
    """Remove transcription artifacts like excessive repetition."""
    # Collapse repeated short phrases (e.g., "Thank you. Thank you. Thank you.")
    text = re.sub(r'((?:Thank you|Merci|OK|Yes|No|Great|Okay)[\.,!]?\s*){4,}',
                  lambda m: m.group(0).split('.')[0].split(',')[0].strip() + '. ', text)
    # Collapse repeated words
    text = re.sub(r'\b(\w+)\b(?:\s+\1\b){3,}', r'\1', text)
    # Clean up extra whitespace
    text = re.sub(r'\s{2,}', ' ', text).strip()
    return text


def main():
    with open("transcript_final.json") as f:
        data = json.load(f)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('Memorial Interview Transcript', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('DEFILE_CAM_45848  |  November 12, 2025  |  Shoah Memorial, Paris')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.space_after = Pt(4)

    # Separator
    sep = doc.add_paragraph()
    run = sep.add_run('_' * 70)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)
    sep.space_after = Pt(12)

    prev_speaker = None
    skipped_empty = 0

    for seg in data["segments"]:
        text = clean_text(seg["text"])
        if not text or len(text) < 3:
            skipped_empty += 1
            continue

        speaker = seg["speaker"]
        ts = format_ts(seg["start"])

        # Add extra spacing when speaker changes
        new_speaker = speaker != prev_speaker

        # Speaker header line (only when speaker changes)
        if new_speaker:
            header = doc.add_paragraph()
            header.space_before = Pt(10) if prev_speaker is not None else Pt(0)
            header.space_after = Pt(2)

            name_run = header.add_run(speaker)
            name_run.bold = True
            name_run.font.size = Pt(11)

            ts_run = header.add_run(f'  [{ts}]')
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(130, 130, 130)
        else:
            # Same speaker continues, just add a subtle timestamp
            header = doc.add_paragraph()
            header.space_before = Pt(2)
            header.space_after = Pt(2)
            ts_run = header.add_run(f'[{ts}]')
            ts_run.font.size = Pt(8)
            ts_run.font.color.rgb = RGBColor(170, 170, 170)

        # Split long text into readable paragraphs
        paragraphs = split_into_paragraphs(text)
        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.space_before = Pt(1)
            p.space_after = Pt(3)
            run = p.add_run(para_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(40, 40, 40)

        prev_speaker = speaker

    doc.save("transcript_final.docx")
    print(f"Formatted {len(data['segments'])} segments ({skipped_empty} empty segments skipped)")
    print("Saved to transcript_final.docx")


if __name__ == "__main__":
    main()
